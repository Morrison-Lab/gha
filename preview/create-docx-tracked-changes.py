#!/usr/bin/env python3
"""Compare rendered DOCX files against the deployed branch and create tracked changes.

Ports `ucdavis/win`'s `.github/scripts/create-docx-tracked-changes.py`
(MIT, copyright 2025 d-morrison), rewritten to this repository's fail-fast bar:

  * `python-docx` is a strict dependency; a missing installation fails fast with a
    clear error rather than silently degrading to copying unchanged files.
  * Every git operation is checked; a missing deployed branch is a stated skip, and
    any network, auth, or repo corruption error raises immediately.
  * No hardcoded `/tmp` scratch paths; base document blobs are read directly from
    git into memory via `cat-file blob` without disk pollution.
  * Step outputs (`docx-status`, `docx-skip-reason`, `docx-tracked-changes-files`,
    `any-docx-changed`) are written to `$GITHUB_OUTPUT` using delimited blocks.

Configuration (all via the environment, set by `preview/action.yml`):

  RENDERED_DIR     Directory holding this run's rendered site. Required.
  DOCX_GLOB        Glob, relative to RENDERED_DIR, selecting the rendered DOCX
                   files to process. Default `chapters/*.docx`.
  DEPLOYED_REMOTE  Git remote holding the published site. Default `origin`.
  DEPLOYED_BRANCH  Branch on that remote. Default `gh-pages`.
  DEPLOYED_SUBDIR  Path prefix, within the deployed branch, at which the
                   published site root lives. Default '' (the branch root).
  REPO_DIR         Git repository to run in. Default `.`.
  GITHUB_OUTPUT    Where step outputs are written, when set.

Outputs:

  docx-status                 `generated` or `skipped`
  docx-skip-reason            Why generation was skipped; empty when generated
  docx-tracked-changes-files  JSON array of generated tracked changes file paths
                              relative to RENDERED_DIR
  any-docx-changed            `true` or `false`
"""

import copy
import difflib
import io
import json
import os
import shutil
import subprocess
import sys
import uuid
from pathlib import Path

try:
    import docx
    from docx import Document
    from docx.opc.part import Part
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.parts.image import ImagePart
except ImportError:
    docx = None
    Document = None
    Part = None
    OxmlElement = None
    qn = None
    ImagePart = None

from _workflow_annotations import annotate


class DocxTrackedChangesError(RuntimeError):
    """A condition that must stop the run rather than generate misleading DOCX files."""


def run_git(args, repo_dir):
    """Run a git command, raising on any non-zero exit."""
    result = subprocess.run(
        ["git", *args],
        cwd=repo_dir,
        capture_output=True,
        text=False,
    )
    if result.returncode != 0:
        stderr = result.stderr.decode("utf-8", "replace").strip()
        raise DocxTrackedChangesError(
            f"git {' '.join(args)} failed in {repo_dir}: {stderr}"
        )
    return result.stdout


def fetch_deployed_branch(repo_dir, remote, branch):
    """Fetch the deployed branch into a local ref, returning the ref name."""
    local_ref = f"refs/remotes/{remote}/{branch}"
    run_git(
        ["fetch", remote, f"+refs/heads/{branch}:{local_ref}"],
        repo_dir,
    )
    return local_ref


def resolve_deployed_ref(repo_dir, remote, branch):
    """Fetch the deployed branch and return a local ref for it, or None if absent."""
    listing = run_git(
        ["ls-remote", "--heads", remote, f"refs/heads/{branch}"], repo_dir
    ).decode("utf-8", "replace")
    if not listing.strip():
        return None

    local_ref = f"refs/gha-preview-base/{branch}"
    run_git(
        [
            "fetch",
            "--no-tags",
            "--depth=1",
            remote,
            f"+refs/heads/{branch}:{local_ref}",
        ],
        repo_dir,
    )
    return local_ref


def published_paths(repo_dir, ref):
    """Every blob path in the deployed tree."""
    listing = run_git(["ls-tree", "-r", "--name-only", "-z", ref], repo_dir)
    return {p for p in listing.decode("utf-8", "replace").split("\0") if p}


def read_published(repo_dir, ref, path):
    """Read a blob's raw bytes directly from git."""
    return run_git(["cat-file", "blob", f"{ref}:{path}"], repo_dir)


def wrap_run_in_ins(run_element, rev_id, author, date):
    """Wrap a Word run element in a w:ins insertion revision tag."""
    parent = run_element.getparent()
    if parent is None:
        return
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(rev_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    parent.insert(parent.index(run_element), ins)
    parent.remove(run_element)
    ins.append(run_element)


def copy_relationships_for_element(element, source_part, target_part):
    """Copy referenced OPC relationships from source_part to target_part and update IDs."""
    if source_part is None or target_part is None:
        return
    rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    attr_prefix = "{" + rel_ns + "}"
    for el in element.iter():
        for attr_name in list(el.attrib.keys()):
            if attr_name.startswith(attr_prefix):
                old_rid = el.attrib[attr_name]
                if hasattr(source_part, "rels") and old_rid in source_part.rels:
                    old_rel = source_part.rels[old_rid]
                    new_rid = None
                    if old_rel.is_external:
                        new_rid = target_part.relate_to(
                            old_rel.target_ref,
                            old_rel.reltype,
                            is_external=True,
                        )
                    elif (
                        (ImagePart is not None and isinstance(getattr(old_rel, "target_part", None), ImagePart))
                        or "image" in getattr(old_rel, "reltype", "")
                    ) and hasattr(target_part, "package") and hasattr(target_part.package, "get_or_add_image_part"):
                        try:
                            image_part = target_part.package.get_or_add_image_part(
                                io.BytesIO(old_rel.target_part.blob)
                            )
                            new_rid = target_part.relate_to(image_part, old_rel.reltype)
                        except Exception:
                            new_rid = None
                    elif hasattr(old_rel, "target_part") and hasattr(target_part, "package") and Part is not None:
                        try:
                            old_target = old_rel.target_part
                            if hasattr(target_part.package, "parts") and any(p.partname == old_target.partname for p in target_part.package.parts):
                                ext = getattr(old_target.partname, "ext", "") or ""
                                template = f"/word/media/part%d.{ext}" if ext else "/word/media/part%d"
                                new_partname = target_part.package.next_partname(template)
                                new_part = Part(new_partname, old_target.content_type, old_target.blob, target_part.package)
                                new_rid = target_part.relate_to(new_part, old_rel.reltype)
                            else:
                                new_rid = target_part.relate_to(old_target, old_rel.reltype)
                        except Exception:
                            new_rid = None
                    if new_rid:
                        el.set(attr_name, new_rid)


def wrap_run_in_del(run_element, rev_id, author, date):
    """Wrap a Word run element in a w:del deletion revision tag, replacing w:t with w:delText."""
    parent = run_element.getparent()
    if parent is None:
        return
    for t_element in run_element.findall(qn("w:t")):
        del_text = OxmlElement("w:delText")
        del_text.set(qn("xml:space"), "preserve")
        del_text.text = t_element.text
        run_element.replace(t_element, del_text)

    del_elem = OxmlElement("w:del")
    del_elem.set(qn("w:id"), str(rev_id))
    del_elem.set(qn("w:author"), author)
    del_elem.set(qn("w:date"), date)
    parent.insert(parent.index(run_element), del_elem)
    parent.remove(run_element)
    del_elem.append(run_element)


def create_docx_with_tracked_changes(
    old_docx_source,
    new_docx_path,
    output_path,
    author="PR Preview",
    # Fixed timestamp (2026-01-01T00:00:00Z) ensures deterministic, reproducible XML
    # so re-rendering unchanged documents does not produce artificial build diffs.
    date="2026-01-01T00:00:00Z",
):
    """Create a DOCX file with tracked changes showing differences against old_docx_source.

    old_docx_source may be bytes, a file path, or None (brand new file).
    """
    if docx is None or Document is None or OxmlElement is None or qn is None:
        raise DocxTrackedChangesError(
            "python-docx is required to generate DOCX tracked changes; "
            "install it with `python3 -m pip install python-docx`"
        )

    # Ensure output directory exists
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Copy new doc to output path as base document
    shutil.copy2(new_docx_path, output_path)

    output_doc = Document(output_path)

    # Enable track revisions in document settings
    settings = output_doc.settings
    settings_element = settings.element
    track_revisions = settings_element.find(qn("w:trackRevisions"))
    if track_revisions is None:
        track_revisions = OxmlElement("w:trackRevisions")
        settings_element.append(track_revisions)

    # Load old and new documents
    if old_docx_source is None:
        old_doc = None
        old_paragraphs = []
    elif isinstance(old_docx_source, bytes):
        old_doc = Document(io.BytesIO(old_docx_source))
        old_paragraphs = [p.text for p in old_doc.paragraphs]
    else:
        old_doc = Document(old_docx_source)
        old_paragraphs = [p.text for p in old_doc.paragraphs]

    new_doc = Document(new_docx_path)
    new_paragraphs = [p.text for p in new_doc.paragraphs]

    matcher = difflib.SequenceMatcher(None, old_paragraphs, new_paragraphs)
    has_changes = False
    rev_id = 0
    original_output_paragraphs = list(output_doc.paragraphs)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag in ("delete", "replace"):
            has_changes = True
            if old_doc is not None:
                for old_idx in range(i1, i2):
                    if old_idx < len(old_doc.paragraphs):
                        old_p = old_doc.paragraphs[old_idx]
                        cloned_p = copy.deepcopy(old_p._element)
                        copy_relationships_for_element(cloned_p, old_doc.part, output_doc.part)
                        r_elements = cloned_p.findall(".//" + qn("w:r"))
                        for r in r_elements:
                            rev_id += 1
                            wrap_run_in_del(r, rev_id, author, date)
                        if j1 < len(original_output_paragraphs):
                            original_output_paragraphs[j1]._element.addprevious(cloned_p)
                        else:
                            sect_pr = output_doc._body._element.find(qn("w:sectPr"))
                            if sect_pr is not None:
                                sect_pr.addprevious(cloned_p)
                            else:
                                output_doc._body._element.append(cloned_p)
        if tag in ("insert", "replace"):
            has_changes = True
            for idx in range(j1, j2):
                if idx < len(original_output_paragraphs):
                    para = original_output_paragraphs[idx]
                    # Find all w:r elements in the paragraph, including those nested in hyperlinks
                    r_elements = para._element.findall(".//" + qn("w:r"))
                    for r_element in r_elements:
                        rev_id += 1
                        wrap_run_in_ins(r_element, rev_id, author, date)

    output_doc.save(output_path)
    return has_changes


def write_outputs(values):
    """Append step outputs in $GITHUB_OUTPUT's delimited form."""
    output_file = os.getenv("GITHUB_OUTPUT")
    if not output_file:
        return
    delimiter = f"gha-eof-{uuid.uuid4().hex}"
    with open(output_file, "a", encoding="utf-8") as handle:
        for key, value in values.items():
            if delimiter in value:
                raise DocxTrackedChangesError(
                    f"output {key!r} contains the generated delimiter; refusing to "
                    "write an output that could be misread"
                )
            handle.write(f"{key}<<{delimiter}\n{value}\n{delimiter}\n")


def generate_tracked_changes(repo_dir, rendered_dir, glob, remote, branch, subdir):
    """Generate tracked changes DOCX files.

    Returns (status, skip_reason, generated_files, any_changed).
    """
    if docx is None or Document is None:
        raise DocxTrackedChangesError(
            "python-docx is required to generate DOCX tracked changes; "
            "install it with `python3 -m pip install python-docx`"
        )

    docx_files = sorted(p for p in rendered_dir.glob(glob) if p.is_file())
    if not docx_files:
        raise DocxTrackedChangesError(
            f"no rendered DOCX files matched {glob!r} under {rendered_dir}; the render "
            "is missing or `docx-tracked-changes-glob` does not match this project"
        )

    ref = resolve_deployed_ref(repo_dir, remote, branch)
    if ref is None:
        reason = (
            f"branch {branch!r} does not exist on remote {remote!r}; nothing has "
            "been deployed yet, so there is no published render to compare against"
        )
        print(annotate("notice", f"Skipping DOCX tracked changes generation: {reason}"))
        return "skipped", reason, [], False

    available = published_paths(repo_dir, ref)
    published_docx_paths = {p for p in available if p.endswith(".docx")}
    if not published_docx_paths:
        reason = (
            f"no published DOCX files found on branch {branch!r}; nothing to compare against"
        )
        print(annotate("notice", f"Skipping DOCX tracked changes generation: {reason}"))
        return "skipped", reason, [], False

    prefix = subdir.strip("/")
    generated_relative_paths = []
    any_changed = False

    for docx_path in docx_files:
        relative = docx_path.relative_to(rendered_dir)
        published_path = f"{prefix}/{relative.as_posix()}" if prefix else relative.as_posix()

        output_filename = f"{docx_path.stem}-tracked-changes.docx"
        output_path = docx_path.parent / output_filename
        output_relative = relative.parent / output_filename

        if published_path not in available:
            # Brand new file in this PR
            has_changes = create_docx_with_tracked_changes(None, docx_path, output_path)
            print(f"  new:       {relative.as_posix()} -> {output_relative.as_posix()}")
        else:
            old_bytes = read_published(repo_dir, ref, published_path)
            has_changes = create_docx_with_tracked_changes(old_bytes, docx_path, output_path)
            if has_changes:
                print(f"  changed:   {relative.as_posix()} -> {output_relative.as_posix()}")
            else:
                print(f"  unchanged: {relative.as_posix()} -> {output_relative.as_posix()}")

        if has_changes:
            any_changed = True
        generated_relative_paths.append(output_relative.as_posix())

    print(
        f"Generated {len(generated_relative_paths)} tracked changes DOCX file(s) "
        f"against {branch!r} (any changed: {any_changed})."
    )
    return "generated", "", generated_relative_paths, any_changed


def main():
    rendered_dir_raw = os.getenv("RENDERED_DIR", "").strip()
    if not rendered_dir_raw:
        raise DocxTrackedChangesError("RENDERED_DIR is required")
    rendered_dir = Path(rendered_dir_raw)
    if not rendered_dir.is_dir():
        raise DocxTrackedChangesError(f"rendered directory {rendered_dir} does not exist")

    status, reason, files, any_changed = generate_tracked_changes(
        repo_dir=os.getenv("REPO_DIR", ".") or ".",
        rendered_dir=rendered_dir,
        glob=os.getenv("DOCX_GLOB", "").strip() or "chapters/*.docx",
        remote=os.getenv("DEPLOYED_REMOTE", "").strip() or "origin",
        branch=os.getenv("DEPLOYED_BRANCH", "").strip() or "gh-pages",
        subdir=os.getenv("DEPLOYED_SUBDIR", ""),
    )

    write_outputs(
        {
            "docx-status": status,
            "docx-skip-reason": reason,
            "docx-tracked-changes-files": json.dumps(files),
            "any-docx-changed": "true" if any_changed else "false",
        }
    )
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except DocxTrackedChangesError as error:
        print(annotate("error", error), file=sys.stderr)
        sys.exit(1)
