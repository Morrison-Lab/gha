"""Tests for `preview/create-docx-tracked-changes.py`.

Ensures DOCX comparison produces valid tracked changes XML using python-docx,
fails fast on missing dependencies or network errors, and states skips cleanly
when no deployed branch or published DOCX files exist.
"""

import io
import json
from pathlib import Path

import pytest

from conftest import read_outputs, write

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    Document = None
    qn = None


def make_docx(paragraphs):
    """Generate a DOCX file with given paragraphs at runtime (no binary fixtures)."""
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def run_generate(docx_generator, monkeypatch, work, rendered_dir, **env):
    """Invoke `main()` through the environment like the composite action does."""
    output_file = rendered_dir.parent / "github-output.txt"
    output_file.write_text("", encoding="utf-8")
    defaults = {
        "REPO_DIR": str(work),
        "RENDERED_DIR": str(rendered_dir),
        "DOCX_GLOB": "chapters/*.docx",
        "DEPLOYED_REMOTE": "origin",
        "DEPLOYED_BRANCH": "gh-pages",
        "DEPLOYED_SUBDIR": "",
        "GITHUB_OUTPUT": str(output_file),
    }
    for key in defaults:
        monkeypatch.delenv(key, raising=False)
    for key, value in {**defaults, **env}.items():
        monkeypatch.setenv(key, str(value))

    docx_generator.main()

    return read_outputs.parse(output_file.read_text(encoding="utf-8"))


def test_missing_deployed_branch_is_a_stated_skip(docx_generator, monkeypatch, repo_factory):
    work = repo_factory(published=None)
    rendered = write(work, "_site/chapters/01.docx", make_docx(["Chapter one"])).parent.parent

    outputs = run_generate(docx_generator, monkeypatch, work, rendered)

    assert outputs["docx-status"] == "skipped"
    assert "gh-pages" in outputs["docx-skip-reason"]
    assert outputs["any-docx-changed"] == "false"
    assert json.loads(outputs["docx-tracked-changes-files"]) == []


def test_no_published_docx_files_is_a_stated_skip(docx_generator, monkeypatch, repo_factory):
    work = repo_factory(published={"chapters/01.html": "<html><body>Chapter one</body></html>"})
    rendered = write(work, "_site/chapters/01.docx", make_docx(["Chapter one"])).parent.parent

    outputs = run_generate(docx_generator, monkeypatch, work, rendered)

    assert outputs["docx-status"] == "skipped"
    assert "no published DOCX" in outputs["docx-skip-reason"]
    assert outputs["any-docx-changed"] == "false"
    assert json.loads(outputs["docx-tracked-changes-files"]) == []


def test_unreachable_remote_is_an_error(docx_generator, monkeypatch, repo_factory, tmp_path):
    work = repo_factory(published={"chapters/01.docx": make_docx(["Chapter one"])})
    rendered = write(work, "_site/chapters/01.docx", make_docx(["Chapter one"])).parent.parent

    with pytest.raises(docx_generator.DocxTrackedChangesError):
        run_generate(
            docx_generator,
            monkeypatch,
            work,
            rendered,
            DEPLOYED_REMOTE=f"file://{tmp_path / 'no-such-repo.git'}",
        )


def test_no_rendered_docx_files_is_an_error(docx_generator, monkeypatch, repo_factory):
    work = repo_factory(published={"chapters/01.docx": make_docx(["Chapter one"])})
    rendered = write(work, "_site/other/empty.txt", "hello").parent.parent

    with pytest.raises(docx_generator.DocxTrackedChangesError, match="no rendered DOCX files matched"):
        run_generate(docx_generator, monkeypatch, work, rendered)


def test_missing_rendered_dir_is_an_error(docx_generator, monkeypatch, repo_factory, tmp_path):
    work = repo_factory(published={"chapters/01.docx": make_docx(["Chapter one"])})
    non_existent = tmp_path / "non_existent_site"

    with pytest.raises(docx_generator.DocxTrackedChangesError, match="does not exist"):
        run_generate(docx_generator, monkeypatch, work, non_existent)


def test_missing_python_docx_is_an_error(docx_generator, monkeypatch, repo_factory):
    work = repo_factory(published={"chapters/01.docx": make_docx(["Chapter one"])})
    rendered = write(work, "_site/chapters/01.docx", make_docx(["Chapter one"])).parent.parent

    monkeypatch.setattr(docx_generator, "docx", None)
    monkeypatch.setattr(docx_generator, "Document", None)

    with pytest.raises(docx_generator.DocxTrackedChangesError, match="python-docx is required"):
        run_generate(docx_generator, monkeypatch, work, rendered)


def test_unchanged_docx_carries_track_revisions_without_spurious_insertions(
    docx_generator, monkeypatch, repo_factory
):
    original_docx = make_docx(["Chapter 1", "Unchanged paragraph content."])
    work = repo_factory(published={"chapters/01.docx": original_docx})
    rendered = write(work, "_site/chapters/01.docx", original_docx).parent.parent

    outputs = run_generate(docx_generator, monkeypatch, work, rendered)

    assert outputs["docx-status"] == "generated"
    assert outputs["any-docx-changed"] == "false"
    assert outputs["docx-skip-reason"] == ""
    assert json.loads(outputs["docx-tracked-changes-files"]) == ["chapters/01-tracked-changes.docx"]

    output_path = work / "_site" / "chapters" / "01-tracked-changes.docx"
    assert output_path.is_file()

    doc = Document(output_path)
    # Track revisions must be enabled in document settings
    assert doc.settings.element.find(qn("w:trackRevisions")) is not None
    # No insertion tags added for unchanged paragraphs
    ins_elements = doc._body._element.findall(".//" + qn("w:ins"))
    assert len(ins_elements) == 0


def test_changed_docx_carries_tracked_changes(docx_generator, monkeypatch, repo_factory):
    old_docx = make_docx(["Chapter 1", "Old paragraph content.", "Conclusion"])
    new_docx = make_docx(["Chapter 1", "Modified paragraph content with edits.", "Conclusion"])
    work = repo_factory(published={"chapters/01.docx": old_docx})
    rendered = write(work, "_site/chapters/01.docx", new_docx).parent.parent

    outputs = run_generate(docx_generator, monkeypatch, work, rendered)

    assert outputs["docx-status"] == "generated"
    assert outputs["any-docx-changed"] == "true"
    assert json.loads(outputs["docx-tracked-changes-files"]) == ["chapters/01-tracked-changes.docx"]

    output_path = work / "_site" / "chapters" / "01-tracked-changes.docx"
    assert output_path.is_file()

    doc = Document(output_path)
    assert doc.settings.element.find(qn("w:trackRevisions")) is not None
    ins_elements = doc._body._element.findall(".//" + qn("w:ins"))
    assert len(ins_elements) > 0
    for ins in ins_elements:
        assert ins.attrib[qn("w:author")] == "PR Preview"
        assert ins.attrib[qn("w:date")] == "2026-01-01T00:00:00Z"


def test_new_docx_file_in_pr_is_generated_with_all_insertions(
    docx_generator, monkeypatch, repo_factory
):
    chapter1_docx = make_docx(["Chapter 1", "Paragraph content."])
    chapter2_docx = make_docx(["Chapter 2 (New)", "Brand new chapter paragraph."])
    work = repo_factory(published={"chapters/01.docx": chapter1_docx})
    write(work, "_site/chapters/01.docx", chapter1_docx)
    rendered = write(work, "_site/chapters/02.docx", chapter2_docx).parent.parent

    outputs = run_generate(docx_generator, monkeypatch, work, rendered)

    assert outputs["docx-status"] == "generated"
    assert outputs["any-docx-changed"] == "true"
    generated_files = json.loads(outputs["docx-tracked-changes-files"])
    assert generated_files == [
        "chapters/01-tracked-changes.docx",
        "chapters/02-tracked-changes.docx",
    ]

    # Chapter 2 should have insertions for its new content
    output_path2 = work / "_site" / "chapters" / "02-tracked-changes.docx"
    assert output_path2.is_file()
    doc2 = Document(output_path2)
    assert doc2.settings.element.find(qn("w:trackRevisions")) is not None
    ins_elements2 = doc2._body._element.findall(".//" + qn("w:ins"))
    assert len(ins_elements2) > 0


def test_custom_deployed_subdir_and_glob(docx_generator, monkeypatch, repo_factory):
    old_docx = make_docx(["Section A", "Old line."])
    new_docx = make_docx(["Section A", "New line."])
    work = repo_factory(published={"nested/site/docs/handout.docx": old_docx})
    rendered = write(work, "_site/docs/handout.docx", new_docx).parent.parent

    outputs = run_generate(
        docx_generator,
        monkeypatch,
        work,
        rendered,
        DOCX_GLOB="docs/*.docx",
        DEPLOYED_SUBDIR="nested/site",
    )

    assert outputs["docx-status"] == "generated"
    assert outputs["any-docx-changed"] == "true"
    assert json.loads(outputs["docx-tracked-changes-files"]) == ["docs/handout-tracked-changes.docx"]
    assert (work / "_site" / "docs" / "handout-tracked-changes.docx").is_file()

def test_deleted_paragraph_reports_any_docx_changed_true(docx_generator, monkeypatch, repo_factory):
    old_docx = make_docx(["Chapter 1", "Paragraph to remove.", "Conclusion"])
    new_docx = make_docx(["Chapter 1", "Conclusion"])
    work = repo_factory(published={"chapters/01.docx": old_docx})
    rendered = write(work, "_site/chapters/01.docx", new_docx).parent.parent

    outputs = run_generate(docx_generator, monkeypatch, work, rendered)

    assert outputs["docx-status"] == "generated"
    assert outputs["any-docx-changed"] == "true"
    assert json.loads(outputs["docx-tracked-changes-files"]) == ["chapters/01-tracked-changes.docx"]
    output_path = work / "_site" / "chapters" / "01-tracked-changes.docx"
    assert output_path.is_file()

    doc = Document(output_path)
    assert doc.settings.element.find(qn("w:trackRevisions")) is not None
    del_elements = doc._body._element.findall(".//" + qn("w:del"))
    assert len(del_elements) > 0
    for d in del_elements:
        assert d.attrib[qn("w:author")] == "PR Preview"
        assert d.attrib[qn("w:date")] == "2026-01-01T00:00:00Z"
    del_texts = [el.text for el in doc._body._element.findall(".//" + qn("w:delText"))]
    assert "Paragraph to remove." in del_texts


def test_replaced_paragraph_carries_both_ins_and_del(docx_generator, monkeypatch, repo_factory):
    old_docx = make_docx(["Chapter 1", "Old content.", "Conclusion"])
    new_docx = make_docx(["Chapter 1", "Replaced content.", "Conclusion"])
    work = repo_factory(published={"chapters/01.docx": old_docx})
    rendered = write(work, "_site/chapters/01.docx", new_docx).parent.parent

    outputs = run_generate(docx_generator, monkeypatch, work, rendered)

    assert outputs["docx-status"] == "generated"
    assert outputs["any-docx-changed"] == "true"
    output_path = work / "_site" / "chapters" / "01-tracked-changes.docx"
    assert output_path.is_file()

    doc = Document(output_path)
    ins_elements = doc._body._element.findall(".//" + qn("w:ins"))
    del_elements = doc._body._element.findall(".//" + qn("w:del"))
    assert len(ins_elements) > 0
    assert len(del_elements) > 0
    del_texts = [el.text for el in doc._body._element.findall(".//" + qn("w:delText"))]
    assert "Old content." in del_texts


def test_trailing_paragraph_deletion(docx_generator, monkeypatch, repo_factory):
    old_docx = make_docx(["Chapter 1", "Ending paragraph to remove."])
    new_docx = make_docx(["Chapter 1"])
    work = repo_factory(published={"chapters/01.docx": old_docx})
    rendered = write(work, "_site/chapters/01.docx", new_docx).parent.parent

    outputs = run_generate(docx_generator, monkeypatch, work, rendered)

    assert outputs["docx-status"] == "generated"
    assert outputs["any-docx-changed"] == "true"
    output_path = work / "_site" / "chapters" / "01-tracked-changes.docx"
    assert output_path.is_file()

    doc = Document(output_path)
    del_elements = doc._body._element.findall(".//" + qn("w:del"))
    assert len(del_elements) > 0
    del_texts = [el.text for el in doc._body._element.findall(".//" + qn("w:delText"))]
    assert "Ending paragraph to remove." in del_texts


def test_paragraph_with_hyperlink_runs_wrapped_properly(docx_generator, monkeypatch, repo_factory):
    from docx.oxml import OxmlElement

    old_docx = make_docx(["Intro", "Old link para"])

    # Create new docx with a hyperlink XML element (runs == [] on paragraph level)
    doc = Document()
    doc.add_paragraph("Intro")
    p = doc.add_paragraph()
    hl = OxmlElement("w:hyperlink")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "https://example.com"
    r.append(t)
    hl.append(r)
    p._element.append(hl)

    buf = io.BytesIO()
    doc.save(buf)
    new_docx_bytes = buf.getvalue()

    work = repo_factory(published={"chapters/01.docx": old_docx})
    rendered = write(work, "_site/chapters/01.docx", new_docx_bytes).parent.parent

    outputs = run_generate(docx_generator, monkeypatch, work, rendered)

    assert outputs["docx-status"] == "generated"
    assert outputs["any-docx-changed"] == "true"
    output_path = work / "_site" / "chapters" / "01-tracked-changes.docx"
    assert output_path.is_file()

    out_doc = Document(output_path)
    ins_elements = out_doc._body._element.findall(".//" + qn("w:ins"))
    assert len(ins_elements) > 0

